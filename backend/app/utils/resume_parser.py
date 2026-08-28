import re
import io
from typing import Dict, Any, List, Optional

try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import docx
except ImportError:
    docx = None

class ResumeParser:
    @staticmethod
    def extract_text_from_pdf(file_bytes: bytes) -> str:
        if pypdf:
            try:
                reader = pypdf.PdfReader(io.BytesIO(file_bytes))
                text = ""
                for page in reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
                if text.strip():
                    return text.strip()
            except Exception as e:
                print(f"[PDF Extraction Warning] {e}")
        try:
            return file_bytes.decode('utf-8', errors='ignore')
        except Exception:
            return ""

    @staticmethod
    def extract_text_from_docx(file_bytes: bytes) -> str:
        if docx:
            try:
                doc = docx.Document(io.BytesIO(file_bytes))
                full_text = [p.text for p in doc.paragraphs if p.text.strip()]
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                full_text.append(cell.text.strip())
                return "\n".join(full_text).strip()
            except Exception as e:
                print(f"[DOCX Extraction Warning] {e}")
        try:
            return file_bytes.decode('utf-8', errors='ignore')
        except Exception:
            return ""

    @staticmethod
    def extract_text_auto(file_bytes: bytes, filename: str = "") -> str:
        ext = filename.lower().split('.')[-1] if '.' in filename else ''
        if ext == 'docx':
            text = ResumeParser.extract_text_from_docx(file_bytes)
            if text.strip():
                return text
        text = ResumeParser.extract_text_from_pdf(file_bytes)
        if text.strip():
            return text
        try:
            return file_bytes.decode('utf-8', errors='ignore')
        except Exception:
            return ""

    @staticmethod
    def extract_sections(text: str) -> Dict[str, str]:
        sections = {
            "summary": "",
            "skills": "",
            "experience": "",
            "projects": "",
            "education": "",
            "certifications": "",
            "achievements": ""
        }

        patterns = {
            "summary": r"(summary|objective|professional summary|about me|profile)",
            "skills": r"(skills|technical skills|technologies|core competencies|tools|skills & competencies)",
            "experience": r"(experience|work experience|employment history|internships|professional experience|work history)",
            "projects": r"(projects|academic projects|key projects|personal projects)",
            "education": r"(education|academic background|qualifications|academic details)",
            "certifications": r"(certifications|certificates|courses|credentials)",
            "achievements": r"(achievements|honors|awards|accomplishments|publications|extracurricular)"
        }

        lines = text.split("\n")
        current_section = None
        buffer = []

        for line in lines:
            trimmed = line.strip()
            if not trimmed:
                continue

            matched = False
            for sec_name, pattern in patterns.items():
                if re.match(r"^#*\s*" + pattern + r"\s*[:\-_]*$", trimmed, re.IGNORECASE):
                    if current_section and buffer:
                        sections[current_section] = "\n".join(buffer).strip()
                    current_section = sec_name
                    buffer = []
                    matched = True
                    break

            if not matched:
                if current_section:
                    buffer.append(trimmed)
                else:
                    if len(sections["summary"]) < 400:
                        sections["summary"] += " " + trimmed

        if current_section and buffer:
            sections[current_section] = "\n".join(buffer).strip()

        return sections

    @staticmethod
    def parse_contact_info(text: str) -> Dict[str, Any]:
        email_match = re.search(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b", text)
        phone_match = re.search(r"(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}", text)
        linkedin_match = re.search(r"(https?://)?(www\.)?linkedin\.com/in/[\w-]+", text, re.IGNORECASE)
        github_match = re.search(r"(https?://)?(www\.)?github\.com/[\w-]+", text, re.IGNORECASE)
        portfolio_match = re.search(r"(https?://)?(www\.)?[\w-]+\.(io|dev|com|me|app)\b", text, re.IGNORECASE)

        # Attempt name extraction from first line
        lines = [line.strip() for line in text.split("\n") if line.strip()]
        name = lines[0] if lines and len(lines[0]) < 40 and not "@" in lines[0] and not "resume" in lines[0].lower() else "Candidate"

        # Attempt location extraction
        loc_match = re.search(r"\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?,\s*(?:[A-Z]{2}|[A-Z][a-z]+))\b", text)
        location = loc_match.group(0) if loc_match else ""

        return {
            "name": name,
            "email": email_match.group(0) if email_match else "",
            "phone": phone_match.group(0) if phone_match else "",
            "linkedin": linkedin_match.group(0) if linkedin_match else "",
            "github": github_match.group(0) if github_match else "",
            "portfolio": portfolio_match.group(0) if portfolio_match and not "linkedin" in portfolio_match.group(0).lower() and not "github" in portfolio_match.group(0).lower() else "",
            "location": location,
            "has_email": bool(email_match),
            "has_phone": bool(phone_match),
            "has_linkedin": bool(linkedin_match),
            "has_github": bool(github_match)
        }

    @staticmethod
    def parse_education(education_text: str, full_text: str) -> List[Dict[str, str]]:
        target_text = education_text if education_text.strip() else full_text
        results = []

        degree_keywords = [
            r"B\.?Tech", r"M\.?Tech", r"Bachelor", r"Master", r"B\.?S\.?", r"M\.?S\.?", r"Ph\.?D",
            r"B\.?E\.?", r"M\.?E\.?", r"Computer Science", r"Information Technology", r"Engineering", r"Diploma", r"Secondary"
        ]
        degree_pattern = r"|".join(degree_keywords)

        lines = [l.strip() for l in target_text.split("\n") if l.strip()]
        for idx, line in enumerate(lines):
            if re.search(degree_pattern, line, re.IGNORECASE):
                year_match = re.search(r"\b(20\d{2}|19\d{2})\b", line) or (re.search(r"\b(20\d{2}|19\d{2})\b", lines[idx+1]) if idx+1 < len(lines) else None)
                inst_match = lines[idx-1] if idx > 0 and len(lines[idx-1]) < 60 else (lines[idx+1] if idx+1 < len(lines) else "University / Institution")
                
                results.append({
                    "degree": line,
                    "institution": inst_match if inst_match != line else "University / College",
                    "year": year_match.group(0) if year_match else "N/A"
                })

        if not results and re.search(degree_pattern, full_text, re.IGNORECASE):
            for match in re.finditer(degree_pattern, full_text, re.IGNORECASE):
                snippet = full_text[max(0, match.start()-30):min(len(full_text), match.end()+50)].replace("\n", " ").strip()
                results.append({
                    "degree": match.group(0),
                    "institution": "University / College",
                    "year": "N/A",
                    "details": snippet
                })

        return results[:4]

    @staticmethod
    def parse_experience(experience_text: str, full_text: str) -> List[Dict[str, Any]]:
        target_text = experience_text if experience_text.strip() else full_text
        results = []

        lines = [l.strip() for l in target_text.split("\n") if l.strip()]
        role_pattern = r"(Engineer|Developer|Intern|Manager|Architect|Analyst|Consultant|Specialist|Lead|Assistant)"
        
        current_exp = None
        for line in lines:
            if re.search(role_pattern, line, re.IGNORECASE) and len(line) < 80:
                if current_exp:
                    results.append(current_exp)
                dates = re.search(r"\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|20\d{2})\b.*?\b(Present|20\d{2})\b", line, re.IGNORECASE)
                current_exp = {
                    "title": line,
                    "company": "Company / Organization",
                    "dates": dates.group(0) if dates else "N/A",
                    "bullets": []
                }
            elif current_exp:
                if line.startswith(('-', '*', '•', '–')):
                    current_exp["bullets"].append(line.lstrip('-*•– ').strip())
                elif len(current_exp["bullets"]) < 5 and len(line) > 20:
                    current_exp["bullets"].append(line)

        if current_exp:
            results.append(current_exp)

        return results[:5]

    @staticmethod
    def parse_projects(projects_text: str, full_text: str) -> List[Dict[str, Any]]:
        target_text = projects_text if projects_text.strip() else full_text
        projects = []

        lines = [l.strip() for l in target_text.split("\n") if l.strip()]
        for idx, line in enumerate(lines):
            if (line.startswith("#") or line.endswith(":") or re.search(r"(System|App|Platform|Portal|Tool|Website|Service|Dashboard|AI|Bot)\b", line, re.IGNORECASE)) and len(line) < 60:
                desc = lines[idx+1] if idx+1 < len(lines) else ""
                techs = re.findall(r"\b(Python|React|FastAPI|Node|MongoDB|PostgreSQL|AWS|Docker|Java|C\+\+|Tailwind|TypeScript)\b", f"{line} {desc}", re.IGNORECASE)
                projects.append({
                    "title": line.strip("#: "),
                    "description": desc,
                    "technologies": list(set(techs))
                })

        return projects[:5]

    @staticmethod
    def parse_skills(skills_text: str, full_text: str) -> Dict[str, List[str]]:
        combined = f"{skills_text} {full_text}".lower()

        tech_catalog = [
            "python", "javascript", "typescript", "java", "c++", "c#", "golang", "rust", "php", "sql",
            "react", "angular", "vue", "next.js", "tailwind", "redux", "html5", "css3",
            "fastapi", "django", "flask", "node.js", "express", "spring boot",
            "mongodb", "postgresql", "mysql", "redis", "elasticsearch",
            "docker", "kubernetes", "aws", "gcp", "azure", "ci/cd", "git", "github"
        ]

        tools_catalog = ["git", "docker", "vscode", "jira", "postman", "figma", "linux", "bash", "jenkins"]
        soft_catalog = ["leadership", "communication", "problem solving", "teamwork", "agile", "scrum", "time management"]

        matched_tech = []
        for t in tech_catalog:
            if re.search(r"\b" + re.escape(t) + r"\b", combined):
                matched_tech.append(t.title() if t not in ["sql", "aws", "gcp", "ci/cd"] else t.upper())

        matched_tools = [t.title() for t in tools_catalog if re.search(r"\b" + re.escape(t) + r"\b", combined)]
        matched_soft = [s.title() for s in soft_catalog if re.search(r"\b" + re.escape(s) + r"\b", combined)]

        return {
            "technical": list(dict.fromkeys(matched_tech)),
            "tools": list(dict.fromkeys(matched_tools)),
            "soft_skills": list(dict.fromkeys(matched_soft))
        }

    @staticmethod
    def parse_certifications(cert_text: str, full_text: str) -> List[Dict[str, str]]:
        target_text = cert_text if cert_text.strip() else full_text
        certs = []

        platforms = ["AWS", "Google", "Microsoft", "Meta", "Coursera", "Udemy", "LinkedIn", "Oracle", "NPTEL", "HackerRank", "LeetCode"]
        lines = [l.strip() for l in target_text.split("\n") if l.strip()]

        for line in lines:
            for p in platforms:
                if re.search(r"\b" + re.escape(p) + r"\b", line, re.IGNORECASE) or "certif" in line.lower():
                    certs.append({
                        "name": line,
                        "issuer": p if p.lower() in line.lower() else "Verified Issuer"
                    })
                    break

        return certs[:5]

    @staticmethod
    def parse_achievements(achieve_text: str, full_text: str) -> List[str]:
        target_text = achieve_text if achieve_text.strip() else full_text
        achievements = []

        keywords = ["awarded", "winner", "ranked", "top", "hackathon", "first place", "certified", "published", "recognized"]
        lines = [l.strip() for l in target_text.split("\n") if l.strip()]

        for line in lines:
            if any(kw in line.lower() for kw in keywords) and len(line) < 120:
                achievements.append(line.lstrip('-*•– '))

        return list(dict.fromkeys(achievements))[:5]

    @classmethod
    def parse_full_structured(cls, text: str) -> Dict[str, Any]:
        raw_sections = cls.extract_sections(text)
        contact = cls.parse_contact_info(text)
        education = cls.parse_education(raw_sections["education"], text)
        experience = cls.parse_experience(raw_sections["experience"], text)
        projects = cls.parse_projects(raw_sections["projects"], text)
        skills = cls.parse_skills(raw_sections["skills"], text)
        certifications = cls.parse_certifications(raw_sections["certifications"], text)
        achievements = cls.parse_achievements(raw_sections["achievements"], text)

        return {
            "contact_info": contact,
            "summary": raw_sections["summary"][:300] if raw_sections["summary"] else "",
            "education": education,
            "experience": experience,
            "projects": projects,
            "skills": skills,
            "certifications": certifications,
            "achievements": achievements
        }
